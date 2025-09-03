from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(parents=True, exist_ok=True)

TITLE = "Stock Predictor Framework – ELI5 Guide"
FILENAME = DOCS / "Stock_Predictor_Framework_Guide.docx"


def add_heading(doc: Document, text: str, level: int = 0):
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def build_document() -> Path:
    doc = Document()

    # Title
    add_heading(doc, TITLE, level=0)
    add_paragraph(doc, " Simple words, small steps.")

    # What is this project
    add_heading(doc, "What is this?", level=1)
    add_paragraph(doc, "It is a helper robot for stocks. It looks at past prices, learns patterns, and makes a tiny guess about tomorrow. Then it tests a simple rule to see if the guesses could make money.")

    # Big pieces
    add_heading(doc, "The big pieces (folders)", level=1)
    add_bullets(doc, [
        "data: gets price data from the internet and saves it on your computer.",
        "features: makes smart numbers (indicators) from prices.",
        "models: teaches a model to guess tomorrow’s return and saves it.",
        "backtest: plays pretend trading with the guesses to see results.",
        "cli: buttons you can push in the terminal to run steps.",
        "streamlit_app.py: a simple app window with buttons.",
    ])

    # Data
    add_heading(doc, "Data – where numbers come from", level=1)
    add_paragraph(doc, "We use yfinance to download stock prices (Open, High, Low, Close, Volume). We save them as files so we don’t have to download again.")

    # Features
    add_heading(doc, "Features – making smart numbers", level=1)
    add_bullets(doc, [
        "Moving averages (SMA/EMA): smooth the price.",
        "RSI/MACD/Bollinger: popular indicators to show trend/momentum/volatility.",
        "Returns: how much price changed over 1, 5, 10 days.",
        "Lags: we shift features by 1 day to avoid cheating with future info.",
        "Target: the thing we guess – next day return (percent change).",
    ])

    # Modeling
    add_heading(doc, "Model – the brain", level=1)
    add_bullets(doc, [
        "RandomForest (default): simple and works on most computers.",
        "XGBoost (optional): sometimes better but needs a macOS library (libomp).",
        "We split the data by time: train, validate, test. No peeking ahead!",
        "We save the model to the models folder so we can use it later.",
    ])

    # Backtest
    add_heading(doc, "Backtest – playing pretend", level=1)
    add_bullets(doc, [
        "Rule: if the model says tomorrow is positive by more than a threshold, we buy; otherwise, we hold cash.",
        "We include simple trade cost.",
        "We compare to Buy & Hold.",
        "We look at results: final money (equity), annual return, Sharpe (risk-adjusted), win rate.",
    ])

    # How to run
    add_heading(doc, "How to run things (easy steps)", level=1)
    add_bullets(doc, [
        "Make a virtual environment and install requirements.",
        "Train: python -m src.stock_predictor.cli train AAPL --start 2018-01-01 --end 2024-12-31 --model-type rf",
        "Backtest: python -m src.stock_predictor.cli run-backtest AAPL --start 2019-01-01 --end 2024-12-31 --threshold 0.0",
        "App: streamlit run streamlit_app.py",
    ])

    # Real-time trading – ELI5
    add_heading(doc, "Real-time trading – ELI5", level=1)
    add_paragraph(doc, "Trading live means: get today’s price, make a fresh guess, decide to buy or not, and place an order at your broker.")
    add_bullets(doc, [
        "Choose a broker API (Alpaca, Interactive Brokers, Robinhood, etc.).",
        "Each broker gives you keys (like passwords) and a small library to send orders.",
        "Write a small script that runs every day after market close:",
    ])
    add_bullets(doc, [
        "1) Download latest data", 
        "2) Make features (same steps as training)",
        "3) Load saved model and predict tomorrow’s return",
        "4) If prediction > threshold, send BUY order; else send SELL/close order.",
    ])

    # Real-time trading – more detail
    add_heading(doc, "Real-time trading – more detail (next steps)", level=2)
    add_bullets(doc, [
        "Scheduler: use cron (Mac/Linux) or launchd to run your script at 3:55pm local time.",
        "Position sizing: start with small size. You can do a fixed dollar amount or percent of cash.",
        "Risk: set max position, stop-loss, and a daily limit.",
        "Logging: save what you did and why (prediction, price, order status).",
        "Paper trading first: most brokers have a paper (fake money) environment. Try there before real money.",
    ])

    # Where to change things (scope for modification)
    add_heading(doc, "Where to change things (scope for modification)", level=1)
    add_bullets(doc, [
        "features/engineering.py: add or remove indicators, change windows.",
        "models/train.py: swap models, tune hyperparameters, add walk-forward CV.",
        "backtest/backtest.py: add position sizing, shorting, slippage, portfolio of many tickers.",
        "data/ingest.py: switch to different data source or intraday interval (like 1h).",
        "cli.py & streamlit_app.py: add new options and visualizations.",
    ])

    # Safety and expectations
    add_heading(doc, "Safety and expectations", level=1)
    add_bullets(doc, [
        "This is not financial advice.",
        "Daily returns are noisy: don’t expect magic. Improve step by step.",
        "Test with paper trading, start small, add risk controls.",
    ])

    doc.save(FILENAME)
    return FILENAME


if __name__ == "__main__":
    path = build_document()
    print(f"Wrote: {path}")
